"""
Diagnostic tool to inspect DOCX structure
"""

from docx import Document
import sys

def diagnose_docx(docx_path):
    doc = Document(docx_path)

    print(f"\n{'='*70}")
    print(f"DOCX Diagnostic Report: {docx_path}")
    print(f"{'='*70}\n")

    # Check paragraphs
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"{'='*70}\n")

    for i, para in enumerate(doc.paragraphs):
        print(f"\n--- Paragraph {i} ---")
        print(f"Text: '{para.text}'")
        print(f"Text length: {len(para.text)}")
        print(f"Text stripped: '{para.text.strip()}'")
        print(f"Stripped length: {len(para.text.strip())}")
        print(f"Style: {para.style.name if para.style else 'None'}")
        print(f"Runs count: {len(para.runs)}")

        for j, run in enumerate(para.runs):
            print(f"\n  Run {j}:")
            print(f"    Text: '{run.text}'")
            print(f"    Font name: {run.font.name}")
            print(f"    Font size: {run.font.size}")
            print(f"    Bold: {run.font.bold}")
            print(f"    Italic: {run.font.italic}")
            print(f"    Underline: {run.font.underline}")
            print(f"    Color: {run.font.color.rgb if run.font.color and run.font.color.rgb else 'None'}")

            # Check style
            if run.style:
                print(f"    Style: {run.style.name}")

            # Check paragraph style font
            if para.style and hasattr(para.style, 'font'):
                print(f"    Para style font: {para.style.font.name}")
                print(f"    Para style size: {para.style.font.size}")

        if i >= 10:  # Limit to first 10 paragraphs
            print(f"\n... (showing first 10 paragraphs only)")
            break

    # Check tables
    print(f"\n{'='*70}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"{'='*70}\n")

    if doc.tables:
        print("Note: This document contains tables!")
        print("The formatscreener.py currently only checks paragraphs.")
        print("Tables need separate handling.\n")

        for t_idx, table in enumerate(doc.tables):
            print(f"\n--- Table {t_idx} ---")
            print(f"Rows: {len(table.rows)}")
            print(f"Columns: {len(table.columns) if table.rows else 0}")

            # Sample first few cells
            for i, row in enumerate(table.rows[:3]):
                for j, cell in enumerate(row.cells[:3]):
                    print(f"\nCell [{i},{j}]:")
                    print(f"  Text: '{cell.text[:50]}...' " if len(cell.text) > 50 else f"  Text: '{cell.text}'")
                    print(f"  Paragraphs in cell: {len(cell.paragraphs)}")

                    # Check first paragraph in cell
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                        if p.runs:
                            r = p.runs[0]
                            print(f"  First run font: {r.font.name}")
                            print(f"  First run size: {r.font.size}")
                            print(f"  First run bold: {r.font.bold}")

            if t_idx >= 2:  # Limit to first 2 tables
                print(f"\n... (showing first 2 tables only)")
                break

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python diagnose-docx.py <path-to-docx>")
        sys.exit(1)

    diagnose_docx(sys.argv[1])
