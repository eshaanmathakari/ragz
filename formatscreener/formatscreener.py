"""
DOCX Format Screener - Style Validation & Scoring Tool

Validates formatting rules for Word documents and provides a compliance score (1-10).

Validation Rules:
- Headings: Calibri, 11pt, Bold, Black
- Body: Calibri, 11pt, No Bold, Black
- Also validates: Font color, Italic status, Underline status

Usage Examples:

  # Quick scoring for simple use cases
  from formatscreener import quick_score

  score = quick_score("resume.docx")
  print(f"Document score: {score}/10")

  # Detailed analysis with categorized violations
  from formatscreener import DocxFormatScreener

  screener = DocxFormatScreener("resume.docx")
  result = screener.score_document()

  print(f"Score: {result['score']}/10")
  print(f"Pass rate: {result['pass_rate']:.1%}")
  print(f"Violations: {result['violations_by_category']}")

  # Agentic AI integration example
  def validate_document(docx_path: str) -> dict:
      screener = DocxFormatScreener(docx_path)
      result = screener.score_document()
      result['meets_standards'] = result['score'] >= 7.0
      return result
"""

import json
from dataclasses import dataclass
from typing import List, Dict, Optional, Any
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.text.run import Run
    from docx.text.paragraph import Paragraph
except ImportError as e:
    raise ImportError(
        "python-docx is required. Install it with: pip install python-docx"
    ) from e


# Public API exports
__all__ = ['DocxFormatScreener', 'FormatRule', 'quick_score']


@dataclass
class FormatObservation:
    """Observed formatting for a run or effective paragraph formatting"""
    font: Optional[str] = None
    size: Optional[float] = None  # in points
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    color: Optional[str] = None  # hex format


@dataclass
class FormatRule:
    """Expected formatting rule"""
    font: str
    size: float
    bold: bool
    italic: bool
    underline: bool
    color: str = "#000000"  # default black


@dataclass
class ValidationResult:
    """Result of validating a single block (paragraph or run)"""
    block_id: str
    block_type: str  # "heading" or "body"
    observed: Dict[str, Any]
    status: str  # "PASS" or "FAIL"
    violations: List[str]
    text_preview: str = ""  # First 50 chars of paragraph text
    violation_details: Dict[str, str] = None  # Human-readable violation messages
    section_name: str = "Unknown Section"  # Actual section heading text

    def __post_init__(self):
        """Initialize violation_details if not provided"""
        if self.violation_details is None:
            self.violation_details = {}


class DocxFormatScreener:
    """Main validator for DOCX formatting"""

    # Define your rules here
    HEADING_RULE = FormatRule(
        font="Calibri",  # change based on client needs
        size=11.0,
        bold=True,
        italic=False,
        underline=False,
        color="#000000"
    )

    BODY_RULE = FormatRule(
        font="Calibri",
        size=11.0,
        bold=False,
        italic=False,
        underline=False,
        color="#000000"
    )

    def __init__(self, docx_path: str, strictness: str = "majority", cluster_consecutive: bool = True):
        """
        Initialize screener

        Args:
            docx_path: Path to the .docx file
            strictness: "strict" (all runs must match) or "majority" (dominant formatting wins)
            cluster_consecutive: If True, group consecutive errors together (default: True)

        Raises:
            FileNotFoundError: If the document doesn't exist
            ValueError: If the file is not a valid DOCX
        """
        self.docx_path = Path(docx_path)
        self.strictness = strictness
        self.cluster_consecutive = cluster_consecutive
        self.current_section_name = None  # Track current section heading for context

        # Validate file exists
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Document not found: {docx_path}")

        # Validate file extension
        if self.docx_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"File must be a .docx or .doc file, got: {self.docx_path.suffix}")

        # Load document with error handling
        try:
            self.document = Document(docx_path)
        except Exception as e:
            raise ValueError(f"Failed to load document (may be corrupted): {str(e)}")

    def _resolve_font_name(self, run: Run, paragraph: Paragraph) -> Optional[str]:
        """
        Resolve effective font name with inheritance fallback

        Priority:
        1. Direct run formatting
        2. Run's character style
        3. Paragraph style
        4. Document defaults / base styles
        """
        # Direct run font
        if run.font.name is not None:
            return run.font.name

        # Character style font
        if run.style and hasattr(run.style, 'font'):
            if run.style.font.name is not None:
                return run.style.font.name

        # Paragraph style font
        if paragraph.style and hasattr(paragraph.style, 'font'):
            if paragraph.style.font.name is not None:
                return paragraph.style.font.name

        # Try to get from paragraph style's base style
        try:
            if paragraph.style and hasattr(paragraph.style, 'base_style'):
                base = paragraph.style.base_style
                if base and hasattr(base, 'font') and base.font.name:
                    return base.font.name
        except:
            pass

        # Could not resolve from explicit styling
        # Try to check document defaults or assume Calibri as Word default
        # For most modern Word documents, None typically means Calibri (Body) theme font
        # However, we should return None and let validation handle it
        # as "unable to determine font"
        return None

    def _resolve_font_size(self, run: Run, paragraph: Paragraph) -> Optional[float]:
        """Resolve effective font size in points"""
        # Direct run size
        if run.font.size is not None:
            return run.font.size.pt

        # Character style size
        if run.style and hasattr(run.style, 'font'):
            if run.style.font.size is not None:
                return run.style.font.size.pt

        # Paragraph style size
        if paragraph.style and hasattr(paragraph.style, 'font'):
            if paragraph.style.font.size is not None:
                return paragraph.style.font.size.pt

        return None

    def _resolve_bold(self, run: Run, paragraph: Paragraph) -> Optional[bool]:
        """Resolve effective bold status"""
        if run.font.bold is not None:
            return run.font.bold

        if run.style and hasattr(run.style, 'font'):
            if run.style.font.bold is not None:
                return run.style.font.bold

        if paragraph.style and hasattr(paragraph.style, 'font'):
            if paragraph.style.font.bold is not None:
                return paragraph.style.font.bold

        return None

    def _resolve_italic(self, run: Run, paragraph: Paragraph) -> Optional[bool]:
        """Resolve effective italic status"""
        if run.font.italic is not None:
            return run.font.italic

        if run.style and hasattr(run.style, 'font'):
            if run.style.font.italic is not None:
                return run.style.font.italic

        if paragraph.style and hasattr(paragraph.style, 'font'):
            if paragraph.style.font.italic is not None:
                return paragraph.style.font.italic

        return None

    def _resolve_underline(self, run: Run, paragraph: Paragraph) -> Optional[bool]:
        """Resolve effective underline status"""
        # Underline is True if it's anything other than None or False
        if run.font.underline is not None:
            return bool(run.font.underline)

        if run.style and hasattr(run.style, 'font'):
            if run.style.font.underline is not None:
                return bool(run.style.font.underline)

        if paragraph.style and hasattr(paragraph.style, 'font'):
            if paragraph.style.font.underline is not None:
                return bool(paragraph.style.font.underline)

        return None

    def _resolve_color(self, run: Run, paragraph: Paragraph) -> Optional[str]:
        """Resolve effective font color as hex"""
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"

        # Character/paragraph style colors are harder to extract reliably
        # For simplicity, we'll default to black if not explicitly set
        return "#000000"

    def _extract_run_formatting(self, run: Run, paragraph: Paragraph) -> FormatObservation:
        """Extract all formatting from a run with inheritance resolution"""
        return FormatObservation(
            font=self._resolve_font_name(run, paragraph),
            size=self._resolve_font_size(run, paragraph),
            bold=self._resolve_bold(run, paragraph),
            italic=self._resolve_italic(run, paragraph),
            underline=self._resolve_underline(run, paragraph),
            color=self._resolve_color(run, paragraph)
        )

    def _get_dominant_formatting(self, paragraph: Paragraph) -> FormatObservation:
        """
        Get dominant formatting across all runs in a paragraph

        In majority mode: pick the most common value for each attribute
        In strict mode: this is used for comparison against all runs
        """
        if not paragraph.runs:
            return FormatObservation()

        # Collect all run observations (ignore whitespace-only runs)
        observations = []
        for run in paragraph.runs:
            if run.text.strip():  # Ignore pure whitespace
                obs = self._extract_run_formatting(run, paragraph)
                observations.append(obs)

        if not observations:
            return FormatObservation()

        # For simplicity in this implementation, use the first substantive run's formatting
        # A more sophisticated approach would count occurrences and pick the mode
        return observations[0]

    def _is_heading(self, paragraph: Paragraph) -> bool:
        """
        Determine if a paragraph is a heading

        Strategy:
        1. Check if style name contains "Heading"
        2. Fallback heuristics:
           - Text is bold AND all uppercase (common pattern: EXPERIENCE, SKILLS)
           - Text is bold AND paragraph text is short (<50 chars) AND followed by content
        """
        style_name = paragraph.style.name if paragraph.style else ""

        # Check style name
        if "Heading" in style_name or "heading" in style_name:
            return True

        # Fallback: heuristic based on formatting and content
        dominant = self._get_dominant_formatting(paragraph)
        text = paragraph.text.strip()

        # Check if text is bold
        if not dominant.bold:
            return False

        # Heuristic: All caps text that's bold (e.g., EXPERIENCE, SKILLS, EDUCATION)
        if text and text.isupper() and len(text) >= 2:
            return True

        # Heuristic: Short bold text (likely a section heading)
        if text and len(text) < 50 and len(text) >= 2:
            return True

        return False

    def _build_observed_dict(self, observed: FormatObservation) -> Dict:
        """
        Build observed dictionary with null values and black color omitted

        Only includes:
        - Non-null font, size, bold
        - Italic/underline only when True (not False/None)
        - Color only when not black (#000000)
        """
        observed_dict = {}

        if observed.font is not None:
            observed_dict["font"] = observed.font
        if observed.size is not None:
            observed_dict["size"] = observed.size
        if observed.bold is not None:
            observed_dict["bold"] = observed.bold
        if observed.italic is not None and observed.italic:
            observed_dict["italic"] = observed.italic
        if observed.underline is not None and observed.underline:
            observed_dict["underline"] = observed.underline
        if observed.color and observed.color.lower() != "#000000":
            observed_dict["color"] = observed.color

        return observed_dict

    def _validate_against_rule(self, observed: FormatObservation, rule: FormatRule) -> tuple[str, List[str], Dict[str, str]]:
        """
        Validate observed formatting against a rule

        Returns: (status, violations_list, violation_details_dict)
        """
        violations = []
        violation_details = {}

        # Font validation
        # Note: None typically means theme font (often Calibri in modern Word docs)
        # We treat None as Calibri for validation purposes
        if observed.font is not None and observed.font != rule.font:
            violations.append("font_mismatch")
            violation_details["font_mismatch"] = f"Expected '{rule.font}' but found '{observed.font}'"
        elif observed.font is None and rule.font != "Calibri":
            # If font is None and we expect something other than Calibri, flag as unknown
            violations.append("font_unknown")
            violation_details["font_unknown"] = f"Expected '{rule.font}' but font could not be determined (likely theme font)"

        if observed.size is not None and abs(observed.size - rule.size) > 0.1:
            violations.append("size_mismatch")
            violation_details["size_mismatch"] = f"Expected {rule.size}pt but found {observed.size}pt"

        if observed.bold is not None and observed.bold != rule.bold:
            violations.append("bold_mismatch")
            violation_details["bold_mismatch"] = f"Expected bold={rule.bold} but found bold={observed.bold}"

        if observed.italic is not None and observed.italic != rule.italic:
            violations.append("italic_mismatch")
            violation_details["italic_mismatch"] = f"Expected italic={rule.italic} but found italic={observed.italic}"

        if observed.underline is not None and observed.underline != rule.underline:
            violations.append("underline_mismatch")
            violation_details["underline_mismatch"] = f"Expected underline={rule.underline} but found underline={observed.underline}"

        if observed.color and observed.color.lower() != rule.color.lower():
            violations.append("color_mismatch")
            violation_details["color_mismatch"] = f"Expected {rule.color} but found {observed.color}"

        status = "PASS" if len(violations) == 0 else "FAIL"
        return status, violations, violation_details

    def validate_paragraph_runs(self, paragraph: Paragraph, para_index: int) -> List[ValidationResult]:
        """
        Validate individual runs (words/clusters) within a paragraph

        Returns a list of ValidationResults, one for each run with formatting issues.
        If all runs conform, returns empty list.
        """
        text = paragraph.text.strip()
        if not text:
            return []

        is_heading = self._is_heading(paragraph)
        block_type = "heading" if is_heading else "body"
        rule = self.HEADING_RULE if is_heading else self.BODY_RULE

        results = []
        run_index = 0

        for run in paragraph.runs:
            run_text = run.text.strip()
            if not run_text:  # Skip empty runs
                continue

            # Extract formatting for this run
            observed = self._extract_run_formatting(run, paragraph)

            # Validate against rule
            status, violations, violation_details = self._validate_against_rule(observed, rule)

            if status == "FAIL":
                # Create result only for failed runs
                observed_dict = self._build_observed_dict(observed)

                result = ValidationResult(
                    block_id=f"p{para_index}_r{run_index}",
                    block_type=block_type,
                    observed=observed_dict,
                    status=status,
                    violations=violations,
                    text_preview=run_text[:50] + "..." if len(run_text) > 50 else run_text,
                    violation_details=violation_details
                )
                results.append(result)

            run_index += 1

        return results

    def validate_paragraph(self, paragraph: Paragraph, para_index: int) -> ValidationResult:
        """
        [DEPRECATED] Validate a single paragraph using dominant formatting

        This method is kept for backward compatibility.
        For new code, use validate_paragraph_runs() for run-level detection.
        """
        # Skip empty paragraphs
        text = paragraph.text.strip()
        if not text:
            return None

        # Determine block type
        is_heading = self._is_heading(paragraph)
        block_type = "heading" if is_heading else "body"
        rule = self.HEADING_RULE if is_heading else self.BODY_RULE

        # Get dominant formatting
        observed = self._get_dominant_formatting(paragraph)

        # Validate
        status, violations, violation_details = self._validate_against_rule(observed, rule)

        # Create text preview (first 50 chars)
        text_preview = text[:50] + "..." if len(text) > 50 else text

        # Build clean observed dict
        observed_dict = self._build_observed_dict(observed)

        # Create result
        result = ValidationResult(
            block_id=f"p_{para_index}",
            block_type=block_type,
            observed=observed_dict,
            status=status,
            violations=violations,
            text_preview=text_preview,
            violation_details=violation_details
        )

        return result

    def validate_document(self) -> List[ValidationResult]:
        """
        Validate entire document at run-level and return results

        Returns a list of ValidationResults, one for each run with formatting issues.
        """
        results = []
        para_counter = 0
        self.current_section_name = None  # Reset at document start

        # Validate regular paragraphs
        for paragraph in self.document.paragraphs:
            # Check if this is a heading and update current section
            if self._is_heading(paragraph):
                heading_text = paragraph.text.strip()
                if heading_text and len(heading_text) < 100:  # Reasonable heading length
                    self.current_section_name = heading_text

            run_results = self.validate_paragraph_runs(paragraph, para_counter)
            # Add section name to each result
            for result in run_results:
                result.section_name = self.current_section_name or "Header"
            results.extend(run_results)
            if paragraph.text.strip():  # Count non-empty paragraphs
                para_counter += 1

        # Validate paragraphs inside tables
        for table_idx, table in enumerate(self.document.tables):
            # Try to determine section from table context
            # Use current section or try first cell
            try:
                first_cell_text = table.rows[0].cells[0].text.strip()
                if first_cell_text and len(first_cell_text) < 50 and first_cell_text.isupper():
                    table_section = first_cell_text
                else:
                    table_section = self.current_section_name or f"Table {table_idx + 1}"
            except (IndexError, AttributeError):
                table_section = self.current_section_name or f"Table {table_idx + 1}"

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for cell_para in cell.paragraphs:
                        # Check if this paragraph is a heading within the table
                        if self._is_heading(cell_para):
                            heading_text = cell_para.text.strip()
                            if heading_text and len(heading_text) < 100:
                                # Update current section for table content
                                self.current_section_name = heading_text
                                table_section = heading_text

                        run_results = self.validate_paragraph_runs(cell_para, para_counter)
                        # Update block_id to include table location and add section name
                        for result in run_results:
                            result.block_id = f"t{table_idx}_r{row_idx}_c{cell_idx}_{result.block_id}"
                            result.section_name = table_section
                        results.extend(run_results)
                        if cell_para.text.strip():  # Count non-empty paragraphs
                            para_counter += 1

        return results

    def _parse_block_id(self, block_id: str) -> dict:
        """
        Parse block_id into components

        Examples:
        - "p5_r0" -> {"para": 5, "run": 0}
        - "t0_r1_c2_p5_r3" -> {"table": 0, "row": 1, "cell": 2, "para": 5, "run": 3}

        Returns:
            Dictionary with parsed components (table, row, cell, para, run)
        """
        parts = block_id.split('_')
        parsed = {}

        for part in parts:
            if part.startswith('t'):
                parsed['table'] = int(part[1:])
            elif part.startswith('r') and 'table' in parsed and 'row' not in parsed:
                parsed['row'] = int(part[1:])
            elif part.startswith('c'):
                parsed['cell'] = int(part[1:])
            elif part.startswith('p'):
                parsed['para'] = int(part[1:])
            elif part.startswith('r'):
                parsed['run'] = int(part[1:])

        return parsed

    def _are_consecutive_blocks(self, block_id1: str, block_id2: str) -> bool:
        """
        Check if two block IDs are consecutive

        Examples:
        - p5 and p6 -> True
        - p5_r0 and p5_r1 -> True (same paragraph, next run)
        - p5_r0 and p6_r0 -> True (next paragraph)
        - p5 and p7 -> False (gap)
        - t0_r0_c0_p5 and t0_r0_c0_p6 -> True

        Returns:
            True if blocks are consecutive, False otherwise
        """
        parsed1 = self._parse_block_id(block_id1)
        parsed2 = self._parse_block_id(block_id2)

        # Must be in same table/row/cell context
        if parsed1.get('table') != parsed2.get('table'):
            return False
        if parsed1.get('row') != parsed2.get('row'):
            return False
        if parsed1.get('cell') != parsed2.get('cell'):
            return False

        # Check paragraph and run indices
        para1 = parsed1.get('para', 0)
        para2 = parsed2.get('para', 0)
        run1 = parsed1.get('run', 0)
        run2 = parsed2.get('run', 0)

        # Same paragraph, consecutive runs
        if para1 == para2 and run2 == run1 + 1:
            return True

        # Consecutive paragraphs (runs can be different)
        if para2 == para1 + 1:
            return True

        return False

    def _should_cluster(self, result1: ValidationResult, result2: ValidationResult) -> bool:
        """
        Determine if two results should be clustered

        Criteria:
        - Same block_type (both heading or both body)
        - Consecutive block IDs

        Returns:
            True if results should be clustered together
        """
        # Must be same block type
        if result1.block_type != result2.block_type:
            return False

        # Must be consecutive
        if not self._are_consecutive_blocks(result1.block_id, result2.block_id):
            return False

        return True

    def _merge_results(self, results: List[ValidationResult]) -> ValidationResult:
        """
        Merge multiple ValidationResults into one clustered result

        - Combine block_ids: "p5~p7" or "p5_r0~p6_r2"
        - Combine text_previews: Join with newline, truncate if >200 chars
        - Union of violations: Deduplicate violation types
        - Merge violation_details: Combine all details
        - Use first result's block_type and observed formatting

        Returns:
            Single merged ValidationResult representing the cluster
        """
        # Get range of block IDs
        first_id = results[0].block_id
        last_id = results[-1].block_id

        # Create range notation
        if len(results) > 1:
            merged_block_id = f"{first_id}~{last_id}"
        else:
            merged_block_id = first_id

        # Combine text previews
        text_parts = [r.text_preview for r in results]
        combined_text = "\n".join(text_parts)
        if len(combined_text) > 200:
            combined_text = combined_text[:197] + "..."

        # Union of violations (deduplicate)
        all_violations = []
        seen = set()
        for r in results:
            for v in r.violations:
                if v not in seen:
                    all_violations.append(v)
                    seen.add(v)

        # Merge violation details
        merged_details = {}
        for r in results:
            for k, v in r.violation_details.items():
                if k not in merged_details:
                    merged_details[k] = v
                else:
                    # Append if different
                    if v not in merged_details[k]:
                        merged_details[k] += f"; {v}"

        # Use first result's formatting as representative
        return ValidationResult(
            block_id=merged_block_id,
            block_type=results[0].block_type,
            observed=results[0].observed,
            status="FAIL",
            violations=all_violations,
            text_preview=combined_text,
            violation_details=merged_details
        )

    def _cluster_consecutive_errors(self, results: List[ValidationResult]) -> List[ValidationResult]:
        """
        Cluster consecutive formatting errors into groups

        Clustering rules:
        - Consecutive block_ids (p5, p6, p7 or p5_r0, p5_r1, etc.)
        - Same block_type (heading or body)

        Returns:
            List of clustered ValidationResults with combined information
        """
        if not results:
            return results

        # Sort by block_id to ensure sequential processing
        sorted_results = sorted(results, key=lambda r: self._parse_block_id(r.block_id).get('para', 0) * 1000 + self._parse_block_id(r.block_id).get('run', 0))

        clusters = []
        current_cluster = [sorted_results[0]]

        for i in range(1, len(sorted_results)):
            prev = sorted_results[i-1]
            curr = sorted_results[i]

            if self._should_cluster(prev, curr):
                current_cluster.append(curr)
            else:
                # Finalize current cluster
                if len(current_cluster) > 1:
                    clusters.append(self._merge_results(current_cluster))
                else:
                    clusters.append(current_cluster[0])

                # Start new cluster
                current_cluster = [curr]

        # Don't forget the last cluster
        if len(current_cluster) > 1:
            clusters.append(self._merge_results(current_cluster))
        else:
            clusters.append(current_cluster[0])

        return clusters

    def calculate_score(self, results: List[ValidationResult]) -> Dict:
        """
        [DEPRECATED] Calculate compliance score from validation results

        This method is kept for backward compatibility.
        For new code, use generate_report() instead which provides
        detailed inconsistency information without scoring.

        Returns:
            Dictionary with score (1-10), pass_rate, and counts
        """
        total = len(results)
        if total == 0:
            return {
                "score": 0.0,
                "pass_rate": 0.0,
                "total_blocks": 0,
                "passed": 0,
                "failed": 0
            }

        passed = sum(1 for r in results if r.status == "PASS")
        pass_rate = passed / total

        # Linear mapping: 0% = 0, 100% = 10
        # Round to 1 decimal place
        score = round(pass_rate * 10, 1)

        return {
            "score": score,
            "pass_rate": pass_rate,
            "total_blocks": total,
            "passed": passed,
            "failed": total - passed
        }

    def categorize_violations(self, results: List[ValidationResult]) -> Dict[str, int]:
        """
        [DEPRECATED] Categorize violations by type

        This method is kept for backward compatibility.
        For new code, use generate_report() instead which provides
        detailed inconsistency information with violation_details.

        Returns:
            Dictionary with counts for each violation category
        """
        categories = {
            "font_errors": 0,
            "size_errors": 0,
            "bold_errors": 0,
            "italic_errors": 0,
            "underline_errors": 0,
            "color_errors": 0
        }

        for result in results:
            for violation in result.violations:
                violation_lower = violation.lower()
                if "font" in violation_lower:
                    categories["font_errors"] += 1
                elif "size" in violation_lower:
                    categories["size_errors"] += 1
                elif "bold" in violation_lower:
                    categories["bold_errors"] += 1
                elif "italic" in violation_lower:
                    categories["italic_errors"] += 1
                elif "underline" in violation_lower:
                    categories["underline_errors"] += 1
                elif "color" in violation_lower:
                    categories["color_errors"] += 1

        return categories

    def score_document(self) -> Dict:
        """
        [DEPRECATED] Main API method: Score document and return structured results

        This method is kept for backward compatibility.
        For new code, use generate_report() instead which provides
        detailed inconsistency information without scoring.

        Returns:
            Dictionary containing:
            - score: float (1-10 scale)
            - pass_rate: float (0.0-1.0)
            - total_blocks: int
            - passed: int
            - failed: int
            - violations_by_category: dict with error counts

        Example:
            screener = DocxFormatScreener("resume.docx")
            result = screener.score_document()
            print(f"Score: {result['score']}/10")
        """
        try:
            # Validate document
            results = self.validate_document()

            # Calculate score
            score_data = self.calculate_score(results)

            # Categorize violations
            violations = self.categorize_violations(results)

            # Combine into final result
            return {
                **score_data,
                "violations_by_category": violations
            }

        except FileNotFoundError:
            raise FileNotFoundError(f"Document not found: {self.docx_path}")
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")

    def _format_comment_string(self, violation_details: Dict[str, str]) -> str:
        """
        Convert violation_details dict to single combined string

        Args:
            violation_details: Dict like {"font_mismatch": "Expected...", "size_mismatch": "Expected..."}

        Returns:
            Single string like "Expected Calibri but found Open Sans; Expected 11.0pt but found 12.0pt"
        """
        if not violation_details:
            return ""

        # Combine all messages with semicolon separator
        messages = list(violation_details.values())
        return "; ".join(messages)

    def _deduplicate_errors(self, inconsistencies: List[ValidationResult]) -> List[ValidationResult]:
        """
        Deduplicate repeated error patterns by adding counts to first occurrence

        Strategy:
        1. Group errors by violation types (same set of violation types)
        2. For each group with >1 occurrence:
           - Keep first occurrence
           - Add count message to its comment: "(N more similar errors found)"

        Returns:
            Deduplicated list with count annotations in comments
        """
        if not inconsistencies:
            return inconsistencies

        # Group by violation signature (sorted violation list)
        from collections import defaultdict
        groups = defaultdict(list)

        for result in inconsistencies:
            # Create signature from sorted violation types
            signature = tuple(sorted(result.violations))
            groups[signature].append(result)

        deduplicated = []

        for signature, results in groups.items():
            # Keep first result
            first_result = results[0]

            # If there are duplicates, add count to comment
            if len(results) > 1:
                additional_count = len(results) - 1
                # Will add count when generating final JSON
                first_result._dedup_count = additional_count

            deduplicated.append(first_result)

        return deduplicated

    def generate_report(self, output_path: Optional[str] = None) -> Dict:
        """
        Generate validation report with formatting inconsistencies only

        Returns a dict with:
        - document: document filename
        - inconsistencies: list with fields: sentence, comment, type of error, Section Name

        The report only includes runs/paragraphs with formatting issues.
        PASS results are not included for cleaner output.

        If cluster_consecutive is True (default), consecutive errors are grouped together.
        """
        results = self.validate_document()

        # Filter to only show failed results (inconsistencies)
        inconsistencies = [r for r in results if r.status == "FAIL"]

        # Apply clustering if enabled
        if self.cluster_consecutive and inconsistencies:
            inconsistencies = self._cluster_consecutive_errors(inconsistencies)

        # Apply deduplication (show first + count for repeated errors)
        inconsistencies = self._deduplicate_errors(inconsistencies)

        # Build report with new field names
        report = {
            "document": str(self.docx_path.name),
            "inconsistencies": []
        }

        # Convert results to new format
        for r in inconsistencies:
            # Convert violation_details dict to single string
            comment = self._format_comment_string(r.violation_details)

            # Add deduplication count if present
            if hasattr(r, '_dedup_count') and r._dedup_count > 0:
                comment += f" ({r._dedup_count} more similar error{'s' if r._dedup_count > 1 else ''} found)"

            # Build new format with only required fields
            result_dict = {
                "sentence": r.text_preview,
                "comment": comment,
                "type of error": r.violations,
                "Section Name": r.section_name
            }

            report["inconsistencies"].append(result_dict)

        # Save to file if requested
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, indent=2, ensure_ascii=False)
            print(f"Report saved to: {output_path}")

        return report

    def print_summary(self):
        """
        [DEPRECATED] Print a human-readable summary (no text content)

        This method is kept for backward compatibility.
        For new code, use generate_report() and format the output as needed.
        """
        result = self.score_document()

        print(f"\n{'='*70}")
        print(f"DOCX Format Screening Report: {self.docx_path.name}")
        print(f"{'='*70}\n")

        print(f"Score: {result['score']}/10")
        print(f"Total blocks analyzed: {result['total_blocks']}")
        print(f"✓ Passed: {result['passed']}")
        print(f"✗ Failed: {result['failed']}")
        print(f"Pass rate: {result['pass_rate']*100:.1f}%\n")

        if result['failed'] > 0:
            print(f"{'='*70}")
            print("VIOLATIONS BY CATEGORY:")
            print(f"{'='*70}\n")

            for category, count in result['violations_by_category'].items():
                if count > 0:
                    category_name = category.replace('_', ' ').title()
                    print(f"  {category_name}: {count}")


def quick_score(docx_path: str) -> float:
    """
    [DEPRECATED] Quick scoring function for simple use cases

    This function is kept for backward compatibility.
    For new code, use DocxFormatScreener().generate_report() instead.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Score from 1-10 based on formatting compliance

    Example:
        score = quick_score("resume.docx")
        print(f"Document score: {score}/10")
    """
    if not Path(docx_path).exists():
        raise FileNotFoundError(f"Document not found: {docx_path}")

    screener = DocxFormatScreener(docx_path)
    result = screener.score_document()
    return result['score']


def main():
    """CLI entry point"""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python formatscreener.py <path-to-docx> [output-json]")
        print("\nExample:")
        print("  python formatscreener.py my-document.docx")
        print("  python formatscreener.py my-document.docx report.json")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not Path(docx_path).exists():
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    # Create screener and run validation
    try:
        screener = DocxFormatScreener(docx_path, strictness="majority")
        report = screener.generate_report(output_path)

        # Print simple summary
        print(f"\n{'='*70}")
        print(f"DOCX Format Screening Report: {Path(docx_path).name}")
        print(f"{'='*70}\n")

        inconsistencies_count = len(report['inconsistencies'])
        if inconsistencies_count == 0:
            print("✓ No formatting inconsistencies found!")
        else:
            print(f"Found {inconsistencies_count} formatting inconsistencies")
            if output_path:
                print(f"\nDetails saved to: {output_path}")

        print(f"\n{'='*70}")
        print("Validation complete!")
        print(f"{'='*70}\n")

    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
