"""Word document parser using python-docx for native parsing."""

import re
from pathlib import Path

import structlog
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table

from rag_assist.ingestion.models import (
    DocumentMetadata,
    DocumentType,
    ParsedDocument,
    SectionContent,
    TextBlock,
)
from rag_assist.ingestion.parsers.base_parser import BaseParser

logger = structlog.get_logger(__name__)


class DOCXParser(BaseParser):
    """Native Word document parser preserving document structure.

    Extracts:
    - Headings with hierarchy
    - Paragraphs
    - Lists
    - Tables
    - Hyperlinks
    """

    supported_types = [DocumentType.DOCX]

    # Heading style patterns
    HEADING_PATTERNS = [
        re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE),
        re.compile(r"^Title$", re.IGNORECASE),
    ]

    def can_parse(self, file_path: str | Path) -> bool:
        """Check if file is a Word document."""
        path = Path(file_path)
        return path.suffix.lower() in [".docx", ".doc"]

    def parse(self, file_path: str | Path, metadata: DocumentMetadata) -> ParsedDocument:
        """Parse Word document.

        Args:
            file_path: Path to DOCX file.
            metadata: Document metadata.

        Returns:
            ParsedDocument with extracted sections.
        """
        self.log_parsing_start(file_path)
        path = Path(file_path)
        sections: list[SectionContent] = []
        errors: list[str] = []

        try:
            doc = Document(str(path))

            # Extract hyperlinks for reference
            hyperlinks = self._extract_all_hyperlinks(doc)

            # Process paragraphs into sections
            current_section = SectionContent()
            page_estimate = 1

            for para in doc.paragraphs:
                try:
                    style_name = para.style.name if para.style else ""
                    heading_level = self._get_heading_level(style_name)

                    if heading_level > 0:
                        # Save previous section if it has content
                        if current_section.heading or current_section.content:
                            sections.append(current_section)

                        # Start new section
                        current_section = SectionContent(
                            heading=self.normalize_text(para.text),
                            heading_level=heading_level,
                            hyperlinks=[
                                h for h in hyperlinks
                                if h.get("paragraph_text") == para.text
                            ],
                        )
                    else:
                        # Add to current section
                        text = self.normalize_text(para.text)
                        if text:
                            block_type = self._get_block_type(para)
                            current_section.content.append(
                                TextBlock(
                                    text=text,
                                    level=self._get_list_level(para),
                                    is_bold=self._is_bold(para),
                                    is_heading=False,
                                    block_type=block_type,
                                )
                            )

                    # Rough page estimation
                    if para.text and len(para.text) > 500:
                        page_estimate += 1

                except Exception as e:
                    error_msg = f"Error parsing paragraph: {str(e)}"
                    errors.append(error_msg)
                    logger.warning(error_msg)

            # Add final section
            if current_section.heading or current_section.content:
                sections.append(current_section)

            # Extract tables
            for table in doc.tables:
                try:
                    table_data = self._extract_table(table)
                    if table_data and sections:
                        # Add table to most recent section
                        sections[-1].tables.append(table_data)
                except Exception as e:
                    error_msg = f"Error parsing table: {str(e)}"
                    errors.append(error_msg)

            metadata.total_pages = page_estimate
            metadata.extraction_method = "native"

        except Exception as e:
            self.log_parsing_error(file_path, e)
            errors.append(f"Failed to open Word document: {str(e)}")

        result = ParsedDocument(
            metadata=metadata,
            content=sections,
            extraction_method="native",
            processing_errors=errors,
        )

        self.log_parsing_complete(file_path, result)
        return result

    def _get_heading_level(self, style_name: str) -> int:
        """Determine heading level from style name.

        Args:
            style_name: Word style name.

        Returns:
            Heading level (1-9) or 0 if not a heading.
        """
        if not style_name:
            return 0

        # Check for Title style
        if style_name.lower() == "title":
            return 1

        # Check for Heading N style
        for pattern in self.HEADING_PATTERNS:
            match = pattern.match(style_name)
            if match and match.groups():
                try:
                    return int(match.group(1))
                except (ValueError, IndexError):
                    pass

        return 0

    def _get_block_type(self, para) -> str:
        """Determine the type of paragraph block.

        Args:
            para: python-docx paragraph object.

        Returns:
            Block type string.
        """
        style_name = para.style.name if para.style else ""
        style_lower = style_name.lower()

        if "list" in style_lower or "bullet" in style_lower:
            return "list_item"
        if "quote" in style_lower:
            return "quote"
        return "paragraph"

    def _get_list_level(self, para) -> int:
        """Get list indentation level.

        Args:
            para: python-docx paragraph object.

        Returns:
            Indentation level (0 for no indentation).
        """
        try:
            if para.paragraph_format.left_indent:
                # Convert to approximate level (each level ~0.5 inch)
                inches = para.paragraph_format.left_indent.inches
                return min(int(inches / 0.5), 5)
        except Exception:
            pass
        return 0

    def _is_bold(self, para) -> bool:
        """Check if paragraph has bold formatting.

        Args:
            para: python-docx paragraph object.

        Returns:
            True if paragraph is bold.
        """
        try:
            for run in para.runs:
                if run.bold:
                    return True
        except Exception:
            pass
        return False

    def _extract_table(self, table: Table) -> list[list[str]]:
        """Extract table data.

        Args:
            table: python-docx table object.

        Returns:
            List of rows, each row is a list of cell texts.
        """
        table_data = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = self.normalize_text(cell.text)
                row_data.append(cell_text)
            table_data.append(row_data)

        return table_data

    def _extract_all_hyperlinks(self, doc: Document) -> list[dict[str, str]]:
        """Extract all hyperlinks from document.

        Args:
            doc: python-docx document object.

        Returns:
            List of hyperlink dictionaries with url and text.
        """
        hyperlinks = []

        try:
            # Get relationship targets
            rels = doc.part.rels
            hyperlink_rels = {
                rel.rId: rel.target_ref
                for rel in rels.values()
                if rel.reltype == RT.HYPERLINK
            }

            # Find hyperlinks in paragraphs
            for para in doc.paragraphs:
                for element in para._element.iter():
                    if element.tag.endswith("hyperlink"):
                        rel_id = element.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        )
                        if rel_id and rel_id in hyperlink_rels:
                            # Get hyperlink text
                            text_parts = []
                            for child in element.iter():
                                if child.text:
                                    text_parts.append(child.text)

                            hyperlinks.append({
                                "url": hyperlink_rels[rel_id],
                                "text": " ".join(text_parts),
                                "paragraph_text": para.text,
                            })

        except Exception as e:
            logger.warning(f"Error extracting hyperlinks: {str(e)}")

        return hyperlinks
